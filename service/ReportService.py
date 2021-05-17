import os

import pdfkit
from docx import Document

from multiprocessing import Pool

from matplotlib import pyplot as plt

from utils import HistogramTools
from utils.LogUtils import LogUtils
from utils.SkinUtils import *
from utils.SkinUtils import SkinUtils
from utils.DistanceUtils import DistanceUtils
from utils.HistogramTools import HistogramTools
from entity.ReportEntity import ReportEntity
from utils.MyDateUtils import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from numpy import doc
from win32com.client import gencache, constants
from core.const_var import *
import cv2


class ReportService:
    @staticmethod
    def wordCreate(face):
        generatePath = OUTPUT_PATH + "\\" + face.name + "\\" + getTodayYearMonthDayHourMinSec()
        if not os.path.isdir(generatePath):
            os.makedirs(generatePath)

        cv2.imwrite(generatePath + "\\" + face.name + "_face.jpg", face.facePart)
        doc = Document()
        # 标题
        doc.add_heading()
        p = doc.add_paragraph()
        title = p.add_run("面容诊断报告")
        title.font.roiName = u'微软雅黑'
        title._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        title.font.size = Pt(24)
        paragraph_format = p.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        def setFontWord(string):
            run = doc.add_paragraph()
            title = run.add_run(string)
            title.font.roiName = u'微软雅黑'
            title._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            paragraph_format = run.paragraph_format
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 内容
        setFontWord("姓名： %s" % face.name)
        setFontWord("性别： %s" % face.gender)
        doc.add_picture(face.imgPath, width=Inches(2))
        doc.save(generatePath + "\\" + face.name + "_report.docx")
        ReportService._word2pdf(generatePath, generatePath + "\\" + face.name + "_report.docx", )

    @staticmethod
    def _word2pdf(wordFolder, wordPath):
        wordfile = wordPath
        pdffile = wordFolder + "\\report.pdf"
        word = gencache.EnsureDispatch('Word.Application')
        doc = word.Documents.Open(wordfile, ReadOnly=1)
        doc.ExportAsFixedFormat(pdffile,
                                constants.wdExportFormatPDF,
                                Item=constants.wdExportDocumentWithMarkup,
                                CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
        word.Quit(constants.wdDoNotSaveChanges)

    @staticmethod
    def long_time_task(roiName, img, colorMode, sampleDict=None):
        """
        多线程执行的方法
        :param roiName:
        :param img:
        :param colorMode:
        :return:
        """
        fig = plt.figure()
        print('Run task %s (%s)...' % (roiName, os.getpid()))
        start = time.time()
        if sampleDict is None:
            sampleDict = ImgUtils.getSampleDict()

        draw = (SkinUtils.skinHistogram(fig, img, colorMode, roiName, sampleDict))
        # draw = (SkinUtils.skinScatter(fig, img, colorMode, roiName, sampleDict))
        end = time.time()
        print('Task %s runs %0.2f seconds.' % (roiName, (end - start)))
        fig.clear()
        return draw

    @staticmethod
    def generateReports(detectedFaces):
        LogUtils.log("reportService", "start generate report from faces...")
        reports = []
        for face in detectedFaces:
            report = ReportEntity()
            report.gender = face.gender
            report.username = face.name
            report.imgPath = face.imgPath
            report.facePart = face.facePart
            report.roiDict = face.landMarkROIDict
            report.roiRGBDict = {}
            report.roiHSVDict = {}
            report.roiYCrCbDict = {}
            report.roiHistograms = {}
            report.roiLabDict = {}
            report.roiColorResults = {}
            # fig = plt.figure(figsize=(5, 5))
            fig = plt.figure()
            LogUtils.log("reportService", "正在获取ROI不同颜色空间报告...")
            items = report.roiDict.items()
            currentProgress = 50
            maxProgress = 100
            step = int((maxProgress - currentProgress) / len(items))

            sstime = time.time()
            result = {}
            p = Pool(len(items))
            sampleDict = ImgUtils.getSampleDict()
            for (name, roi) in items:
                nameCN = FACIAL_LANDMARKS_NAME_DICT[name]
                LogUtils.log("ReportService", "开启线程为" + nameCN + "绘图中... ")
                result[name] = p.apply_async(ReportService.thread_roi_process,
                                             args=(name, roi, sampleDict, currentProgress))

                # p.apply_async(ReportService.long_time_task,     args=(data, "plot" + str(i),))

            p.close()
            p.join()

            for (name, roi) in items:
                nameCN = FACIAL_LANDMARKS_NAME_DICT[name]
                LogUtils.log("ReportService", "线程正在为" + nameCN + "绘图中... ", progress=currentProgress)
                drawItems = result[name].get()
                report.roiRGBDict[name] = drawItems['rgb']
                report.roiHSVDict[name] = drawItems['hsv']
                report.roiYCrCbDict[name] = drawItems['ycrcb']
                report.roiLabDict[name] = drawItems['lab']
                res = drawItems['res']
                report.roiHistograms[name] = res['hist']
                report.roiColorResults[name] = res['result']
                currentProgress += step
                LogUtils.log("ReportService", "线程为" + nameCN + "绘图完毕！", progress=currentProgress)

            report.drawImg = face.drawImg
            eetime = time.time()
            LogUtils.log("reportService", "获取ROI不同颜色空间报告完成！总用时:" + str(eetime - sstime))
            # report.faceColor = SkinUtils.roiTotalColorDetect(report.rois)
            # report.skinResult = SkinUtils.getResultByColor(report.rois)
            reports.append(report)

        LogUtils.log("reportService", "reports generate finished!")
        return reports

    @staticmethod
    def thread_roi_process(name, roi, sampleDict, currentProgress):
        stime = time.time()
        nameCN = FACIAL_LANDMARKS_NAME_DICT[name]
        LogUtils.log("ReportService", "正在绘制" + nameCN + "的颜色空间分布曲线")
        nameCN = FACIAL_LANDMARKS_NAME_DICT[roi.roiName]
        predictImg = roi.img

        rgbDict = ReportService.long_time_task(name, predictImg, COLOR_MODE_RGB, sampleDict)

        hsvDict = ReportService.long_time_task(name, predictImg, COLOR_MODE_HSV, sampleDict)

        LabDict = ReportService.long_time_task(name, predictImg, COLOR_MODE_Lab, sampleDict)

        YCrCbDict = ReportService.long_time_task(name, predictImg, COLOR_MODE_YCrCb, sampleDict)

        LogUtils.log("ReportService", "查询" + nameCN + "的肤色中...")

        res = ReportService.threadGetDistance(roi.img, name, sampleDict)

        etime = time.time()
        LogUtils.log("ReportService", "线程为" + nameCN + "绘图完毕!, 用时:" + str(etime - stime), progress=currentProgress)
        return {
            'rgb': rgbDict,
            'lab': LabDict,
            'hsv': hsvDict,
            'ycrcb': YCrCbDict,
            'res': res
        }

    @staticmethod
    def threadGetDistance(img, name, sampleDict):
        fig = plt.figure()
        roiHistograms, roiColorResults, = HistogramTools.getDistanceByDifferentColorSpace(fig, img, name, sampleDict)
        d = {
            'hist': roiHistograms,
            'result': roiColorResults,
        }
        return d
