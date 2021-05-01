'''
    本程序实现功能是人脸检测，并分割人脸区域，然后进行面部健康诊断
'''
import cv2 as cv
import math
import os
import numpy as np
import copy
import pdfkit
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from matplotlib import pyplot as plt
# from tongueDiagnose import TongueDiagnose
from const_var import *
from win32com.client import gencache, constants

'''
    基于图片的人脸检测
    scaleFactor 指定在每个图像所放缩图像大小减少了多少
    minNeighbors 指定每个候选矩形应该保留多少个邻居
    minsize 最小可能的对象大小
    maxsize 最大可能的对象大小
'''


def faceDetect(img, flag):
    cascPath = OPENCV_CASCADE_PATH + "\\haarcascade_frontalface_default.xml"
    eyePath = OPENCV_CASCADE_PATH + "\\haarcascade_eye.xml"
    smilePath = OPENCV_CASCADE_PATH + "\\haarcascade_smile.xml"

    faceCascade = cv.CascadeClassifier(cascPath)
    eyeCascade = cv.CascadeClassifier(eyePath)
    smileCascade = cv.CascadeClassifier(smilePath)

    faces = faceCascade.detectMultiScale(img,
                                         scaleFactor=1.1,
                                         minNeighbors=5,
                                         minSize=(100, 100),
                                         flags=cv.CASCADE_SCALE_IMAGE)
    for (x, y, w, h) in faces:
        # Draw rectangle around the face
        # print(x, y, w, h)
        cv.rectangle(img, (x, y), (x + w, y + h), (255, 255, 255), 3)
        face_part = img[y:y + h, x:x + w]
        cv.imwrite(OUTPUT_PATH + '\\facePart.jpg', face_part)

        # 肤色检测
        L0, A0, B0, ind, color, gloss, gloss_index = face_color(face_part, face_part)
        # print(L0, A0, B0, ind, color, gloss)

        # 斑点检测
        face_with_spots = spotDefect(face_part)

        img[y:y + h, x:x + w] = face_with_spots

        '''
        使用YCrCb方法进行进行皮肤部分抠图
        '''
        # 把图像转换到YUV色域
        ycrcb = cv.cvtColor(face_part, cv.COLOR_BGR2YCrCb)
        # 图像分割，分别获取y, Cr, Cb 通道图像
        (y, Cr, Cb) = cv.split(ycrcb)
        # 高斯滤波， cr 是带滤波的源图像数据, (5， 5)是值窗口的大小， 0是根据窗口大小来计算高斯函数的标准差
        cr1 = cv.GaussianBlur(Cr, (5, 5), 0)
        # _, skin1 = cv.threshold(cr1, 0, 255, cv.THRESH_BINARY + cv.THRESH_OTSU)
        _, skin1 = cv.threshold(cr1, 0, 255, cv.THRESH_TOZERO + cv.THRESH_OTSU)

        # 生成结果图片
        cv.imwrite(OUTPUT_PATH + '\\DiagnoseResult.jpg', img)

    if flag == 0: return img
    if flag == 1: return color, gloss, img

    '''
    将肤色在YUV色域的图像和圈出人脸的图像进行展示
    '''
    # plt.figure(figsize = (12,8))

    # ax1 = plt.subplot(2, 3, 1)
    # plt.imshow(face_part)
    # ax1.set_title('face')

    # ax2 = plt.subplot(2, 3, 2)
    # plt.imshow(cr1)
    # ax2.set_title('cr')

    # ax3 = plt.subplot(2, 3, 3)
    # plt.imshow(skin1)
    # ax3.set_title('skin')

    # 原图需要从BGR色域转为RGB色域
    # ax4 = plt.subplot2grid((2, 3), (1, 0), colspan = 3)
    # plt.imshow(cv.cvtColor(img, cv.COLOR_BGR2RGB))
    # ax4.set_title('image')

    # plt.imshow(img, cmap = 'gray')    # 显示为灰度图像
    # plt.show()

    # return color, gloss, img
    # return img


# 斑点检测
def spotDefect(img):
    # spot_det = cv.imread("facePart.jpg", -1)
    spot_det = copy.deepcopy(img)
    spot_detector = cv.SimpleBlobDetector_create()
    spotPoints = spot_detector.detect(spot_det)
    face_with_spots = cv.drawKeypoints(spot_det, spotPoints, np.array([]), (255, 0, 0),
                                       cv.DRAW_MATCHES_FLAGS_DRAW_RICH_KEYPOINTS)
    # cv.imshow('spot', face_with_spots)
    # cv.waitKey(0)
    return face_with_spots


'''
    视频中人脸检测
'''


def videoFacedetect():
    cascPath = OPENCV_CASCADE_PATH + "\\haarcascade_frontalface_default.xml"
    eyePath = OPENCV_CASCADE_PATH + "\\haarcascade_eye.xml"
    smilePath = OPENCV_CASCADE_PATH + "\\haarcascade_smile.xml"

    faceCascade = cv.CascadeClassifier(cascPath)
    eyeCascade = cv.CascadeClassifier(eyePath)
    smileCascade = cv.CascadeClassifier(smilePath)

    video_capture = cv.VideoCapture(0)
    while True:
        # Capture frame-by-frame
        ret, frame = video_capture.read()
        img = cv.cvtColor(frame, cv.COLOR_BGR2GRAY)
        faces = faceCascade.detectMultiScale(img,
                                             scaleFactor=1.1,
                                             minNeighbors=5,
                                             minSize=(30, 30),
                                             flags=cv.CASCADE_SCALE_IMAGE)
        for (x, y, w, h) in faces:
            if w > 50:
                cv.rectangle(frame, (x, y), (x + w, y + h), (255, 0, 0), 3)
                roi_img = img[y:y + h, x:x + w]
                roi_color = frame[y:y + h, x:x + w]

                smile = smileCascade.detectMultiScale(roi_img,
                                                      scaleFactor=1.16,
                                                      minNeighbors=35,
                                                      minSize=(25, 25),
                                                      flags=cv.CASCADE_SCALE_IMAGE)
                for (sx, sy, sw, sh) in smile:
                    cv.rectangle(roi_color, (sh, sy), (sx + sw, sy + sh), (255, 0, 0))
                    cv.putText(frame, 'Smile', (x + sx, y + sy), 1, 1, (0, 255, 0), 1)

                eyes = eyeCascade.detectMultiScale(roi_img)
                for (ex, ey, ew, eh) in eyes:
                    cv.rectangle(roi_color, (ex, ey), (ex + ew, ey + eh), (0, 255, 0), 2)
                    cv.putText(frame, 'Eye', (x + ex, y + ey), 1, 1, (0, 255, 0), 1)

        cv.putText(frame, 'Number of faces: ' + str(len(faces)), (40, 40), 1, 1, (255, 0, 0), 2)
        # display the resulting frame
        cv.imshow('Video', frame)

        if cv.waitKey(1) & 0xFF == ord('q'):
            break
    video_capture.release()
    cv.destroyAllWindows()


def face_color(faceskinimage, faceblock):
    """
    面部颜色检测
    :param faceskinimage:
    :param faceblock:
    :return:
    """
    img_lab = cv.cvtColor(faceskinimage, cv.COLOR_BGR2Lab)
    L_value, A_value, B_value = cv.split(img_lab)
    L0 = int(round(np.mean(L_value)))
    A0 = int(round(np.mean(A_value)))
    B0 = int(round(np.mean(B_value)))

    fcs = {"青": [194, 87, 116], "赤": [110, 188, 166], "黄": [218, 135, 150], "白": [245, 129, 134],
           "黑": [96, 142, 143], "正常": [232, 134, 142]}

    if L0 <= 100:
        face_color = "黑"
        ind = -1
    if L0 > 100:
        df1 = ((fcs["青"][1] - A0) ** 2 + (fcs["青"][2] - B0) ** 2) ** 0.5
        df2 = ((fcs["赤"][1] - A0) ** 2 + (fcs["赤"][2] - B0) ** 2) ** 0.5
        df3 = ((fcs["黄"][1] - A0) ** 2 + (fcs["黄"][2] - B0) ** 2) ** 0.5
        df4 = ((fcs["白"][1] - A0) ** 2 + (fcs["白"][2] - B0) ** 2) ** 0.5
        df5 = ((fcs["正常"][1] - A0) ** 2 + (fcs["正常"][2] - B0) ** 2) ** 0.5
        df = [df1, df2, df3, df4, df5]
        ind = df.index(min(df))
        if ind == 0: face_color = "青"
        if ind == 1: face_color = "赤"
        if ind == 2: face_color = "黄"
        if ind == 3: face_color = "白"
        if ind == 4: face_color = "正常"
    gloss_index_temp = round(1.3 * face_gloss_index(faceblock), 2)
    gloss_index = gloss_index_temp if gloss_index_temp <= 0.98 else 0.98
    if gloss_index >= 0.7:
        gloss = "有光泽"
    elif 0.4 < gloss_index < 0.7:
        gloss = "少光泽"
    else:
        gloss = "无光泽"
    return L0, A0, B0, ind, face_color, gloss, gloss_index


def face_gloss_index(faceskinimage):
    img = cv.resize(faceskinimage, (10, 10), interpolation=cv.INTER_AREA)
    c = 80
    sum = 0
    B, G, R = cv.split(img)
    rows, cols = R.shape[0], R.shape[1]
    for i in range(rows):
        for j in range(cols):
            sum = sum + math.exp((-1) * ((i * i + j * j) / (c * c)))
    lamb = 1.0 / sum
    Fxy = np.zeros((rows, cols))
    for i in range(rows):
        for j in range(cols):
            Fxy[i][j] = lamb * math.exp((-1) * ((i * i + j * j) / (c * c)))
    return round((face_gloss_index_pre(B, Fxy) + face_gloss_index_pre(G, Fxy) + face_gloss_index_pre(R, Fxy)) / 3, 2)


def face_gloss_index_pre(single_channel_img, Fxy):
    img = single_channel_img / 255
    rows, cols = img.shape[0], img.shape[1]
    VecB = np.zeros((rows * 2 - 1, cols * 2 - 1))
    answer = np.zeros((rows * 2 - 1, cols * 2 - 1))
    for i in range(rows * 2 - 1):
        for j in range(cols * 2 - 1):
            temp = 0
            for m in range(rows):
                for n in range(cols):
                    if 0 <= (i - m) < rows and 0 <= (j - n) < cols:
                        temp = temp + Fxy[m][n] * img[i - m][j - n]
            VecB[i][j] = math.log10(temp)
    sum = 0
    for i in range(rows):
        for j in range(cols):
            # answer[i][j] = math.pow(10, (math.log10(img[i][j])-VecB[i][j]))
            answer[i][j] = math.log10(img[i][j]) - VecB[i][j]
            sum = sum + answer[i][j]
    gloss_index = sum / (rows * cols)
    return gloss_index


def CreateDetectResults(face_color, face_gloss):
    if face_color == '正常': resultOfSkin = "“漂亮的皮囊”get，请继续保持！"
    if face_color == '白': resultOfSkin = "“漂亮的皮囊”get，请继续保持！"
    if face_color == '黄': resultOfSkin = "“漂亮的皮囊”get，请继续保持！"
    if face_color == '赤': resultOfSkin = "“漂亮的皮囊”get，请继续保持！"
    if face_color == '青': resultOfSkin = "“漂亮的皮囊”get，请继续保持！"
    if face_color == '黑': resultOfSkin = "最近皮肤有点差哦，请注意多休息，可以使用一些护肤品！"

    if face_gloss == '有光泽': resultOfGloss = '如果不是皮肤太油，那就是皮肤太好，羡慕~'
    if face_gloss == '无光泽': resultOfGloss = '面色红润有光泽，你不想要吗，用一些保湿的护肤品吧！'
    if face_gloss == '少光泽': resultOfGloss = '皮肤有点光泽，用一点护肤品就更好了。'

    return resultOfSkin, resultOfGloss


def txt2pdf(txt_file, pdf_file):
    # 将字符串生成pdf文件
    config = pdfkit.configuration(wkhtmltopdf=r"D:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe")
    options = {'encoding': 'utf-8'}
    pdfkit.from_file(txt_file, pdf_file, configuration=config, options=options)


def pdfCreate(filename, name, sex, face_color, face_gloss, resultOfSkin, resultOfGloss):
    with open(filename, 'w') as f:
        # f.write("姓名: %s <br>" %'孙悦')
        # f.write("性别: %s <br>" %'男')
        # f.write("面部诊断结果: %s <br>" %'正常')
        f.write("姓名: %s <br>" % name)
        f.write("性别: %s <br>" % sex)
        f.write("您的面部诊断结果: <br>")
        f.write("&emsp;&emsp; 肤色诊断结果: %s <br>" % face_color)
        f.write("&emsp;&emsp; 皮肤光泽诊断结果: %s <br>" % face_gloss)
        f.write("诊断结果分析: <br>&emsp;&emsp; %s <br>" % resultOfSkin)
        f.write("&emsp;&emsp; %s <br>" % resultOfGloss)

    with open(filename, 'r') as f:
        txt2pdf(f, 'DetectResults.pdf')


def wordCreate(name, gender, faceColor, faceGloss, SkinResults, GlossResults, image):
    if gender == 1: Gender = '男'
    if gender == 0: Gender = '女'
    doc = Document()

    # 标题
    doc.add_heading()
    p = doc.add_paragraph()
    title = p.add_run("面容诊断报告")
    title.font.name = u'微软雅黑'
    title._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    title.font.size = Pt(24)
    paragraph_format = p.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def setFontWord(string):
        run = doc.add_paragraph()
        title = run.add_run(string)
        title.font.name = u'微软雅黑'
        title._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        paragraph_format = run.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 内容
    setFontWord("姓名： %s" % name)
    setFontWord("性别： %s" % Gender)
    setFontWord("面容健康诊断：")
    setFontWord("    肤色状态：%s" % faceColor)
    setFontWord("    光泽状态：%s" % faceGloss)
    doc.add_picture(image, width=Inches(2))
    setFontWord("面容健康小贴士：")
    setFontWord("    (1)%s" % SkinResults)
    setFontWord("    (2)%s" % GlossResults)
    doc.save(OUTPUT_PATH + '\\FaceDiagnoseResults.docx')


def word2pdf(wordfile, pdffile):
    word = gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(wordfile, ReadOnly=1)
    doc.ExportAsFixedFormat(pdffile,
                            constants.wdExportFormatPDF,
                            Item=constants.wdExportDocumentWithMarkup,
                            CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
    word.Quit(constants.wdDoNotSaveChanges)


if __name__ == "__main__":
    # 输入检测人员信息
    filename = "faceDetectResults.txt"
    wordfile = "D:\GitDoc\FaceHealthDetect\FaceDiagnoseResults.docx"
    pdffile = "D:\GitDoc\FaceHealthDetect" + "\FaceDiagnoseResults.pdf"
    name = "孙悦"
    gender = 1

    # 加载图片
    # img = cv.imread('testYao.jpg', cv.IMREAD_COLOR)
    # img = cv.imread('testOfC.jpg', cv.IMREAD_COLOR)
    img = cv.imread('selfieOfSun.jpg', cv.IMREAD_COLOR)
    # img = cv.imread('InkedselfieOfSun.jpg', cv.IMREAD_COLOR)

    # 显示原始图片
    # cv.imshow('origin', img)
    # cv.waitKey(0)

    # 进行人脸检测
    faceColor, faceGloss, img = faceDetect(img, 1)

    # 根据人脸检测情况和人员信息，生成诊断结果
    SkinResults, GlossResults = CreateDetectResults(faceColor, faceGloss)
    # pdfCreate(filename, name, sex, faceColor, faceGloss, SkinResults, GlossResults)
    image = "DiagnoseResult.jpg"
    wordCreate(name, gender, faceColor, faceGloss, SkinResults, GlossResults, image)
    word2pdf(wordfile, pdffile)
